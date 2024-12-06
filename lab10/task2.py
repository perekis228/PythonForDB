from docx import Document, shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document('lab10.docx')

document.add_picture('Wallpaper.jpg', width=shared.Cm(15))
par = document.add_paragraph()
par.add_run('Рисунок 1')
par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.save('lab10.docx')