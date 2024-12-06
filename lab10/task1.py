from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def grey_bold(cell, text):
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="C0C0C0"/>'.format(nsdecls('w'))))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def center(cell, text):
    paragraph = cell.add_paragraph()
    paragraph.add_run(text)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#Создание
document = Document()

#1 абзац
document.add_paragraph('В микроконтроллерах ATmega, используемых на платформах Arduino существует три вида памяти:')

#1 айтем списка
document.add_paragraph('Флеш-память: используется для хранения скетчей.', style='ListBullet')

#2 айтем списка
par = document.add_paragraph('ОЗУ(', style='ListBullet')
run = par.add_run('SRAM')
run.bold = True

par.add_run(' — ')
run = par.add_run('static random access memory')
run.italic = True

par.add_run(', статическая оперативная память с произвольным доступом): используется для хранения и работы переменных.')

#3 айтем списка
document.add_paragraph('EEPROM (энергонезависимая память): используется для хранения постоянной информации.', style='ListBullet')

#2 абзац
document.add_paragraph('Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.')

#Таблица
table = document.add_table(rows=4, cols=5)

for i, text in enumerate(['', 'ATmega168', 'ATmega328', 'ATmega1280', 'ATmega2560']):
    grey_bold(table.cell(0, i), text)

for i, text in enumerate(['Flash(1 кБ flash-памяти занят загрузчиком)', 'SRAM', 'EEPROM']):
    grey_bold(table.cell(i+1, 0), text)

for i, text in enumerate(['16 Кбайт', '32 Кбайт', '128 Кбайт', '256 Кбайт']):
    center(table.cell(1, i+1), text)

for i, text in enumerate(['1 Кбайт', '2 Кбайт', '8 Кбайт', '8 Кбайт']):
    center(table.cell(2, i+1), text)

for i, text in enumerate(['512 байт', '1024 байта', '4 Кбайт', '4 Кбайт']):
    center(table.cell(3, i+1), text)

#3 абзац
par = document.add_paragraph()
run = par.add_run('Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25 С. Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM.')
run.italic = True

#Сейв
document.save('lab10.docx')